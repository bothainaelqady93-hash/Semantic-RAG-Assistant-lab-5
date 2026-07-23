"""
01_documents.py
================
Document loading stage of the RAG pipeline.

Responsibility
--------------
Load raw source documents (from a local folder or from files uploaded through
the Streamlit UI) and normalize them into a single, consistent schema that
every later stage of the pipeline (preprocessing, chunking, embedding,
storage, retrieval) can rely on.

This mirrors the document schema used throughout Lab 8 and Lab 9, where every
document carries:
    - document_id   : a stable integer identifier
    - title         : human readable name of the document
    - doc_type      : a coarse category (policy, procedure, faq, guide, ...)
    - source        : file name or path the document came from
    - effective_date: an ISO date string used to resolve current vs. outdated
                       conflicts during retrieval
    - is_current    : whether this document is the up-to-date version
    - text          : the raw extracted text

Supported input formats: .txt, .md, .pdf, .docx

Design decision
----------------
The labs use hand-written Python dictionaries as the corpus. A real
application needs to load arbitrary user-provided files, so this module adds
a thin, well-tested loading layer on top of that same schema instead of
inventing a new one. This keeps every downstream stage (chunking, embedding,
prompting) identical to what was taught in the labs.
"""

from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


@dataclass
class LoadedDocument:
    """A single normalized source document."""

    document_id: int
    title: str
    doc_type: str
    source: str
    effective_date: str
    is_current: bool
    text: str
    extra_metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Return the document as a plain dictionary (used by later stages)."""
        payload = {
            "document_id": self.document_id,
            "title": self.title,
            "doc_type": self.doc_type,
            "source": self.source,
            "effective_date": self.effective_date,
            "is_current": self.is_current,
            "text": self.text,
        }
        payload.update(self.extra_metadata)
        return payload


def _infer_doc_type(filename: str) -> str:
    """
    Infer a coarse document type from the filename.

    This is a lightweight heuristic (not NLP-based) that looks for common
    keywords such as "policy", "faq", or "guide". If nothing matches, the
    document is labeled "document" as a safe default.

    Args:
        filename: Original file name, e.g. "tuition_refund_policy.txt".

    Returns:
        A short, human-readable document type string.
    """
    name = filename.lower()
    keyword_to_type = {
        "policy": "policy",
        "faq": "faq",
        "guide": "guide",
        "procedure": "procedure",
        "notice": "notice",
        "manual": "manual",
        "handbook": "handbook",
        "checklist": "checklist",
    }
    for keyword, doc_type in keyword_to_type.items():
        if keyword in name:
            return doc_type
    return "document"


def _title_from_filename(filename: str) -> str:
    """
    Build a readable title from a file name.

    Example: "tuition_refund_policy.txt" -> "Tuition Refund Policy"

    Args:
        filename: The original file name including extension.

    Returns:
        A human-friendly title string.
    """
    stem = os.path.splitext(os.path.basename(filename))[0]
    stem = re.sub(r"[_\-]+", " ", stem).strip()
    return stem.title() if stem else "Untitled Document"


def _read_txt_bytes(raw_bytes: bytes) -> str:
    """Decode raw bytes from a .txt/.md file into text, tolerating bad bytes."""
    try:
        return raw_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return raw_bytes.decode("utf-8", errors="ignore")


def _read_pdf_bytes(raw_bytes: bytes) -> str:
    """
    Extract text from PDF bytes using pypdf.

    Args:
        raw_bytes: Raw PDF file content.

    Returns:
        Extracted plain text, with pages joined by double newlines.

    Raises:
        RuntimeError: If the PDF cannot be parsed.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise RuntimeError(
            "pypdf is required to read PDF files. Install it with `pip install pypdf`."
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(raw_bytes))
        pages_text = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages_text)
    except Exception as exc:  # noqa: BLE001 - surface a clear, wrapped error
        raise RuntimeError(f"Failed to parse PDF content: {exc}") from exc


def _read_docx_bytes(raw_bytes: bytes) -> str:
    """
    Extract text from a .docx file's raw bytes using python-docx.

    Args:
        raw_bytes: Raw DOCX file content.

    Returns:
        Extracted plain text, with paragraphs joined by newlines.

    Raises:
        RuntimeError: If the DOCX cannot be parsed.
    """
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise RuntimeError(
            "python-docx is required to read DOCX files. Install it with `pip install python-docx`."
        ) from exc

    try:
        document = docx.Document(io.BytesIO(raw_bytes))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"Failed to parse DOCX content: {exc}") from exc


def extract_text_from_bytes(raw_bytes: bytes, filename: str) -> str:
    """
    Dispatch raw file bytes to the correct extractor based on file extension.

    Args:
        raw_bytes: Raw file content.
        filename: Original filename, used only to determine the extension.

    Returns:
        Extracted plain text.

    Raises:
        ValueError: If the file extension is not supported.
    """
    extension = os.path.splitext(filename)[1].lower()

    if extension in (".txt", ".md"):
        return _read_txt_bytes(raw_bytes)
    if extension == ".pdf":
        return _read_pdf_bytes(raw_bytes)
    if extension == ".docx":
        return _read_docx_bytes(raw_bytes)

    raise ValueError(
        f"Unsupported file extension '{extension}'. "
        f"Supported extensions are: {sorted(SUPPORTED_EXTENSIONS)}"
    )


def load_document_from_bytes(
    raw_bytes: bytes,
    filename: str,
    document_id: int,
    effective_date: Optional[str] = None,
    is_current: bool = True,
) -> LoadedDocument:
    """
    Build a single LoadedDocument from in-memory file bytes.

    This is the code path used by the Streamlit file uploader, where files
    never touch disk.

    Args:
        raw_bytes: Raw file content.
        filename: Original file name (used for title/type inference).
        document_id: Stable integer id to assign to this document.
        effective_date: ISO date string; defaults to today if not provided.
        is_current: Whether this document should be treated as up to date.

    Returns:
        A populated LoadedDocument instance.

    Raises:
        ValueError: If the text extracted from the file is empty.
    """
    text = extract_text_from_bytes(raw_bytes, filename).strip()
    if not text:
        raise ValueError(f"No extractable text found in '{filename}'.")

    return LoadedDocument(
        document_id=document_id,
        title=_title_from_filename(filename),
        doc_type=_infer_doc_type(filename),
        source=filename,
        effective_date=effective_date or datetime.now(timezone.utc).strftime("%Y-%m-%d"),
        is_current=is_current,
        text=text,
    )


def load_documents_from_uploads(uploaded_files: List[Any]) -> List[Dict[str, Any]]:
    """
    Load documents from a list of Streamlit `UploadedFile` objects.

    Args:
        uploaded_files: Objects as returned by `st.file_uploader(..., accept_multiple_files=True)`.
                         Each object must expose `.name` and `.getvalue()`.

    Returns:
        A list of document dictionaries (see LoadedDocument.to_dict()).

    Raises:
        ValueError: If no valid documents could be loaded.
    """
    documents: List[Dict[str, Any]] = []
    errors: List[str] = []

    for document_id, uploaded_file in enumerate(uploaded_files):
        try:
            raw_bytes = uploaded_file.getvalue()
            loaded = load_document_from_bytes(
                raw_bytes=raw_bytes,
                filename=uploaded_file.name,
                document_id=document_id,
            )
            documents.append(loaded.to_dict())
        except (ValueError, RuntimeError) as exc:
            errors.append(f"{uploaded_file.name}: {exc}")

    if not documents:
        error_summary = "; ".join(errors) if errors else "no files provided"
        raise ValueError(f"Could not load any documents. Details: {error_summary}")

    return documents


def load_documents_from_directory(directory: str) -> List[Dict[str, Any]]:
    """
    Load every supported file from a local directory.

    Args:
        directory: Path to a folder containing source documents.

    Returns:
        A list of document dictionaries, sorted by file name for reproducibility.

    Raises:
        FileNotFoundError: If the directory does not exist.
        ValueError: If no supported files were found.
    """
    if not os.path.isdir(directory):
        raise FileNotFoundError(f"Directory not found: {directory}")

    file_names = sorted(
        f for f in os.listdir(directory)
        if os.path.splitext(f)[1].lower() in SUPPORTED_EXTENSIONS
    )

    if not file_names:
        raise ValueError(
            f"No supported files ({sorted(SUPPORTED_EXTENSIONS)}) found in '{directory}'."
        )

    documents: List[Dict[str, Any]] = []
    for document_id, file_name in enumerate(file_names):
        file_path = os.path.join(directory, file_name)
        with open(file_path, "rb") as file_handle:
            raw_bytes = file_handle.read()

        mtime = os.path.getmtime(file_path)
        effective_date = datetime.fromtimestamp(mtime, tz=timezone.utc).strftime("%Y-%m-%d")

        loaded = load_document_from_bytes(
            raw_bytes=raw_bytes,
            filename=file_name,
            document_id=document_id,
            effective_date=effective_date,
        )
        documents.append(loaded.to_dict())

    return documents


if __name__ == "__main__":
    # Simple manual smoke test: load the bundled sample documents.
    sample_dir = os.path.join(os.path.dirname(__file__), "data", "raw")
    try:
        docs = load_documents_from_directory(sample_dir)
        print(f"Loaded {len(docs)} document(s) from '{sample_dir}':")
        for doc in docs:
            print(f"  - [{doc['document_id']}] {doc['title']} ({doc['doc_type']}, {len(doc['text'])} chars)")
    except (FileNotFoundError, ValueError) as exc:
        print(f"Could not load sample documents: {exc}")
